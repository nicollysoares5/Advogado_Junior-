# app.py — IA do Advogado Júnior (versão final e corrigida)
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import textwrap
import re

# ----------- CONFIGURAÇÃO DE PÁGINA -----------
st.set_page_config(
    page_title="IA do Advogado Júnior ⚖️",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------- ESTILO PERSONALIZADO -----------
st.markdown("""
<style>
/* Fundo geral */
body { background-color: #f7f9fb; }
/* Sidebar */
[data-testid="stSidebar"]: # "{ background: linear-gradient(180deg, #0a1e3d 0%, #173a6d 100%); color: white; }"
[data-testid="stSidebar"]: # "* { color: white !important; font-family: 'Georgia', serif; }"
/* Título principal */
h1 { color: #0a1e3d; font-family: 'Georgia', serif; text-align: center; font-weight: bold; }
/* Subtítulos e seções */
h2, h3 { color: #173a6d; font-family: 'Georgia', serif; }
/* Botões */
div.stButton > button { background-color: #173a6d; color: white; border-radius: 8px; padding: 0.5rem 1rem; border: none; font-weight: bold; }
div.stButton > button:hover { background-color: #204d94; color: #fff; }
/* Cards */
div.block-container { padding-top: 1rem; }
.stCard { background-color: white; border-radius: 12px; box-shadow: 0 3px 6px rgba(0,0,0,0.1); padding: 1rem 1.5rem; margin-bottom: 1.2rem; }
/* Footer */
footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ----------- FUNÇÕES AUXILIARES -----------
def summarize_text(text, n_sentences=3):
    text = text.strip()
    if not text: return "Nenhum texto fornecido para resumo."
    sents = re.split(r'(?<=[.!?])\s+', text)
    return " ".join(sents[:n_sentences])

def improve_text(text):
    if not text.strip(): return "Nenhum texto fornecido para melhoria."
    text = re.sub(r'\s+', ' ', text.strip())
    return f"Considerando o exposto, {text[0].upper() + text[1:]}"

def transform_to_petition_from_text(text, petitioner, respondent, city):
    facts = summarize_text(text, 8)
    return f"""EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA COMARCA DE {city.upper()}

{petitioner}, por seu advogado, vem propor a presente AÇÃO DE RESPONSABILIDADE CIVIL em face de {respondent}, pelos fatos e fundamentos jurídicos a seguir expostos:

DOS FATOS
{facts}

DO DIREITO
Aplica-se o dever geral de reparação conforme o Código Civil e a jurisprudência consolidada.

DOS PEDIDOS
Requer:
a) Citação do réu;
b) Condenação por danos morais e materiais;
c) Produção de provas;
d) Condenação em custas e honorários.

{city}, ____ de __________ de 20__.

__________________________________
Advogado/OAB
"""

def create_docx_from_text(title, text):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_from_text(title, text):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, title)
    y -= 20
    c.setFont("Helvetica", 10)
    for line in textwrap.wrap(text, 100):
        c.drawString(50, y, line)
        y -= 12
        if y < 50:
            c.showPage()
            y = height - 50
    c.save()
    buf.seek(0)
    return buf

# ----------- FUNÇÃO PARA CARREGAR O DICIONÁRIO DO CSV -----------
@st.cache_data
def load_dictionary():
    """Carrega o dicionário jurídico a partir de um arquivo CSV."""
    try:
        df = pd.read_csv("dicionario_juridico.csv")
        if "termo" in df.columns and "definicao" in df.columns:
            return dict(zip(df["termo"].str.lower(), df["definicao"]))
        else:
            st.error("O arquivo 'dicionario_juridico.csv' deve conter as colunas 'termo' e 'definicao'.")
            return {}
    except FileNotFoundError:
        st.error("Arquivo 'dicionario_juridico.csv' não encontrado. Certifique-se de que ele está na mesma pasta que o app.py.")
        return {}

# ----------- CABEÇALHO COM LOGO -----------
st.markdown("""
    <div style="background: linear-gradient(90deg, #0a1e3d, #173a6d); padding: 1rem; text-align:center; border-radius:10px; margin-bottom:20px;">
        <img src="https://upload.wikimedia.org/wikipedia/commons/9/9a/Scale_of_justice.png" width="80" style="margin-bottom:10px;">
        <h1 style="color:white; font-family:'Georgia';">IA do Advogado Júnior ⚖️</h1>
    </div>
""", unsafe_allow_html=True )

# ----------- SIDEBAR -----------
st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/9/9a/Scale_of_justice.png", width=90 )
st.sidebar.title("IA do Advogado Júnior ⚖️")
st.sidebar.markdown("---")
menu = st.sidebar.radio("📚 Navegação", ["Assistente Jurídico", "Gerador de Petições", "Dicionário Jurídico"])
st.sidebar.markdown("---")
st.sidebar.caption("Desenvolvido por Nicolly Soares Mota e Maria Eduarda Bustamante Fontoura 💼")

# ----------- CONTEÚDO PRINCIPAL -----------
if menu == "Assistente Jurídico":
    st.markdown("<h1>📝 Assistente Jurídico</h1>", unsafe_allow_html=True)
    st.markdown("Cole o texto jurídico abaixo e escolha a ação desejada.")
    texto = st.text_area("Texto Jurídico", height=250, placeholder="Cole aqui o texto da decisão, ementa ou petição...")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🔍 Resumir"):
            st.markdown("<div class='stCard'><h3>Resumo Gerado</h3>", unsafe_allow_html=True)
            result = summarize_text(texto)
            st.write(result)
            st.markdown("</div>", unsafe_allow_html=True)
    with col2:
        if st.button("✍️ Melhorar Texto"):
            st.markdown("<div class='stCard'><h3>Texto Aperfeiçoado</h3>", unsafe_allow_html=True)
            result = improve_text(texto)
            st.write(result)
            st.markdown("</div>", unsafe_allow_html=True)
    with col3:
        if st.button("📑 Transformar em Petição"):
            petitioner = st.text_input("Autor", "Fulano de Tal")
            respondent = st.text_input("Réu", "Empresa X")
            city = st.text_input("Cidade/UF", "Brasília/DF")
            petition = transform_to_petition_from_text(texto, petitioner, respondent, city)
            st.markdown("<div class='stCard'><h3>Petição Gerada</h3>", unsafe_allow_html=True)
            st.text(petition)
            st.markdown("</div>", unsafe_allow_html=True)
            docx_b = create_docx_from_text("Petição", petition)
            pdf_b = create_pdf_from_text("Petição", petition)
            st.download_button("📥 Baixar .docx", data=docx_b, file_name="peticao.docx")
            st.download_button("📄 Baixar .pdf", data=pdf_b, file_name="peticao.pdf")

elif menu == "Gerador de Petições":
    st.markdown("<h1>📄 Gerador de Petições</h1>", unsafe_allow_html=True)
    tipo = st.selectbox("Tipo de ação", ["Dano moral", "Cobrança", "Mandado de Segurança", "Habeas Corpus", "Outro"])
    autor = st.text_input("Autor", "Fulano de Tal")
    advogado = st.text_input("Advogado/OAB", "Dr. Exemplo - OAB/UF 0000")
    reu = st.text_input("Réu", "Empresa X")
    fatos = st.text_area("Fatos", height=100)
    pedidos = st.text_area("Pedidos", height=100)
    cidade = st.text_input("Cidade", "Brasília/DF")
    if st.button("🧾 Gerar Petição"):
        texto = f"""EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DA COMARCA DE {cidade.upper()}

{autor}, por seu advogado ({advogado}), vem propor a presente AÇÃO DE {tipo.upper()} em face de {reu}, pelos fatos e fundamentos:

DOS FATOS
{fatos}

DOS PEDIDOS
{pedidos}

{cidade}, ____ de __________ de 20__.

__________________________________
{advogado}
"""
        st.markdown("<div class='stCard'><h3>Petição Gerada</h3>", unsafe_allow_html=True)
        st.text(texto)
        st.markdown("</div>", unsafe_allow_html=True)
        docx_b = create_docx_from_text("Petição", texto)
        pdf_b = create_pdf_from_text("Petição", texto)
        st.download_button("📥 Baixar .docx", data=docx_b, file_name="peticao_gerada.docx")
        st.download_button("📄 Baixar .pdf", data=pdf_b, file_name="peticao_gerada.pdf")

elif menu == "Dicionário Jurídico":
    st.markdown("<h1>📚 Dicionário Jurídico</h1>", unsafe_allow_html=True)
    
    # Carrega o dicionário do arquivo CSV
    defs = load_dictionary()

    termo = st.text_input("Digite o termo jurídico:", placeholder="Ex.: acórdão, posse, usucapião")

    if st.button("Buscar definição"):
        termo_limpo = termo.strip().lower()
        if not defs:
            # A mensagem de erro já foi exibida pela função load_dictionary
            pass
        elif termo_limpo == "":
            st.info("Digite um termo para buscar.")
        elif termo_limpo in defs:
            st.success(f"**{termo.capitalize()}:** {defs[termo_limpo]}")
        else:
            # Tenta encontrar termos similares
            similares = [k for k in defs.keys() if termo_limpo in k]
            if similares:
                st.info(f"Termo não encontrado. Mostrando resultado para **'{similares[0].capitalize()}'**:")
                st.success(f"**{similares[0].capitalize()}:** {defs[similares[0]]}")
            else:
                st.warning(f"O termo **'{termo}'** não foi encontrado no dicionário.")
    
    st.markdown("---")
    # Citação da fonte dos dados
    st.caption("Fonte dos dados: Termos e definições compilados a partir do Glossário do Judiciário (Conselho Nacional de Justiça - CNJ) e do Código Civil brasileiro.")
